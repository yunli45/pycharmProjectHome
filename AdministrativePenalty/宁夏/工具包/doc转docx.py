import win32com
from win32com.client import Dispatch
# doc 转为docx，传入：转入文件路径+文件名 +转换后保存路径+文件名
docName = "E:\Python\PyCharm\project\AdministrativePenalty\宁夏\中华人民共和国产品质量法.doc"
docxName = "E:\Python\PyCharm\project\AdministrativePenalty\宁夏\中华人民共和国产品质量法.docx"
def docConvertDocx(docName,docxName):
    # 导入库
    word_app = win32com.client.Dispatch('Word.Application')
    # 调用word程序
    word_app.Visible = 0
    word_app.DisplayAlerts = 0
    # 不在前台显示文档及错误，在实际使用阶段可以全部关闭，提高运行速度，但是在调试时打开还是用处挺大的，可以对操作是否实现自己的需求进行直观的判断，比如说我们选中的内容是否已经高亮等等。
    doc = word_app.Documents.Open(u'%s'%(docName))
    doc.SaveAs(r"%s"%(docxName), 16)
    # 文档路径，暂时未尝试路径能否直接使用变量解决
    # 路径前的u还是必要的，尤其是路径中存在中文时，否则会报错
    #使用参数16表示将doc转换成docx
    """
    # title_list = []
# for lenth in range(0,len_tables):
# #这里的len_tables是在其他函数中已经定义的变量，其为读取的整个文档表格个数。
#     table_index = lenth + 1
#     #个人感觉win32com库操作word时也用到了很多vba的功能，例如很多代码就和vba比较相似，但是又有一些不同，所以在操作时报错报到怀疑人生。
#     #在这里doc.Range().Tables(table_index)中，表格的序号不再是从0开始，而是从1开始。
#     for i in range(0,5):
#     #循环获取指定表格前的第一段文字，一般第一段文字就是表格标题，绝大多数的文档排版后表格前一般会有一段空行，但一般也不会太多。
#         x = doc.Range().Tables(table_index).Range.Start
#         #获取表格开始的位置，即在整个word文档中的位置
#         f = doc.Paragraphs(doc.Range(0,x).Paragraphs.Count - i).Range()
#         #获取最后一段的内容，其原理为先计算从文档开头到表格开始时总共有几段，然后从最后一段开始向前数
#         if f.replace(" ","").replace("\n","").replace("\r","") != "":
#         #去除换行符及空格影响，并且判断是否为空行。
#             title_list.append(f)
#             #将获取的表格标题放到列表中，由于列表有顺序，之后取用也就方便多了。
#             break
    """
    doc.Close()
    # 关闭文档
    word_app.Quit()
    # 关闭word程序

    #读取word内容
    #import docx
    # doc = docx.Document("D:\most.docx")
    # data = doc.paragraphs[0].text
    # print(data)
    return docxName
    print("转换完成，保存的文件路径是："+str(docxName)+",已经返回了该路径给函数")
# 传入docx文件的绝对路径
def readDocx(docxName):

    from docx import Document
    # 创建文档对象
    document = Document(r'%s'%(docxName))
    # 读取文档中所有的段落列表
    ps = document.paragraphs
    # 每个段落有两个属性：style和text
    ps_detail = [(x.text, x.style.name) for x in ps]
    with open('out.txt', 'w+') as fout:
        fout.write('')
    # 读取段落并写入一个文件
    with open('out.txt', 'a+') as fout:
        for p in ps_detail:
            docxCon = "<p>"+p[0]+"</p>" '\t'
    print(docxCon)
    # 读取文档中的所有段落的列表
    # tables = document.tables
    # 遍历table，并将所有单元格内容写入文件中
    # with open('out.tmp', 'a+') as fout:
    #     for table in tables:
    #         for row in table.rows:
    #             for cell in row.cells:
    #                 fout.write(cell.text + '\t')
    #             fout.write('\n')
# docConvertDocx(docName,docxName)
readDocx('E:\Python\PyCharm\project\AdministrativePenalty\宁夏\ggsddss.docx')